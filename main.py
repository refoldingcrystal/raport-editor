import sys
import os
import re
import pathlib
from PyQt6.QtWidgets import (QApplication, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QTreeWidget, QTreeWidgetItem, 
                             QFileDialog, QLabel, QStackedWidget, QListWidget, 
                             QTreeWidgetItemIterator, QLineEdit, QMessageBox)
from PyQt6.QtCore import Qt

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnexpectedEndOfFileError

def convert(path, pattern):
    p = pathlib.Path(path)
    parent_folder = p.parent.name 
    clean_folder = re.sub(r'^\d+-', '', parent_folder)
    filename_no_ext = p.stem
    filename_no_num = re.sub(r'\s\(\d+\)$', '', filename_no_ext)
    result = pattern.replace('%f', clean_folder)
    result = result.replace('%F', parent_folder)
    result = result.replace('%p', filename_no_num)
    result = result.replace('%P', filename_no_ext)
    return result


class FileApp(QWidget):
    def full_report(self):
        doc = Document(self.manual_file_path)
        
        new_section = doc.add_section()
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

        pattern = self.name_suffix_input.text()
        valid_files = [f for f in self.selected_files if os.path.exists(f) and os.path.getsize(f) > 0]

        for i in range(0, len(valid_files), 2):
            table = doc.add_table(rows=2, cols=2)
            try:
                table.style = 'Table Grid'
            except:
                pass

            path_left = valid_files[i]
            self._add_entry_to_table(table, 0, path_left, pattern)

            if i + 1 < len(valid_files):
                path_right = valid_files[i+1]
                self._add_entry_to_table(table, 1, path_right, pattern)
            
            doc.add_page_break()

        doc.save(self.final_save_path)
        self.show_confirmation_popup()

    def _add_entry_to_table(self, table, col_idx, path, pattern):
        try:
            new_name = convert(path, pattern)
            
            header_cell = table.cell(0, col_idx)
            header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_txt = header_cell.paragraphs[0].add_run(new_name)
            run_txt.bold = True
            
            img_cell = table.cell(1, col_idx)
            para = img_cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = para.add_run()
            run_img.add_picture(path, width=Inches(4))
            
        except (UnexpectedEndOfFileError, Exception) as e:
            table.cell(1, col_idx).text = f"Error loading: {path}"

    def __init__(self):
        super().__init__()
        self.selected_files = []
        self.manual_file_path = ""
        self.final_save_path = ""
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Raport Editor')
        self.setGeometry(300, 300, 900, 600)
        
        self.stack = QStackedWidget()
        self.screen1 = self.create_selection_screen()
        self.screen2 = self.create_preview_screen()
        self.screen3 = self.create_final_entry_screen()
        
        self.stack.addWidget(self.screen1)
        self.stack.addWidget(self.screen2)
        self.stack.addWidget(self.screen3)
        
        main_layout = QVBoxLayout()
        main_layout.addWidget(self.stack)
        self.setLayout(main_layout)
        self.validate_screen1()

    def create_selection_screen(self):
        page = QWidget()
        layout = QVBoxLayout()
        self.btn_open = QPushButton('Wybierz katalog główny')
        self.btn_open.clicked.connect(self.pick_folder)
        self.tree = QTreeWidget()
        self.tree.setHeaderLabel("Zaznacz foldery i pliki")
        
        self.tree.itemChanged.connect(self.handle_check_state)
        
        self.btn_next = QPushButton('Dalej')
        self.btn_next.clicked.connect(self.go_to_preview)
        
        layout.addWidget(self.btn_open)
        layout.addWidget(self.tree)
        layout.addWidget(self.btn_next)
        page.setLayout(layout)
        return page

    def create_preview_screen(self):
        page = QWidget()
        layout = QVBoxLayout()
        
        input_layout = QHBoxLayout()
        input_layout.addWidget(QLabel("Opis zdjęcia"))
        self.name_suffix_input = QLineEdit("%f %p")
        self.name_suffix_input.textChanged.connect(self.refresh_preview_list)
        input_layout.addWidget(self.name_suffix_input)
        layout.addLayout(input_layout)

        panes = QHBoxLayout()
        self.list_original = QListWidget()
        self.list_converted = QListWidget()
        panes.addWidget(self.list_original)
        panes.addWidget(self.list_converted)
        
        self.list_original.verticalScrollBar().valueChanged.connect(self.list_converted.verticalScrollBar().setValue)
        self.list_converted.verticalScrollBar().valueChanged.connect(self.list_original.verticalScrollBar().setValue)
        
        nav_layout = QHBoxLayout()
        self.btn_back_1 = QPushButton('Wstecz')
        self.btn_back_1.clicked.connect(lambda: self.stack.setCurrentIndex(0))
        self.btn_go_final = QPushButton('Dalej')
        self.btn_go_final.clicked.connect(lambda: self.stack.setCurrentIndex(2))
        
        nav_layout.addWidget(self.btn_back_1)
        nav_layout.addWidget(self.btn_go_final)
        layout.addLayout(panes)
        layout.addLayout(nav_layout)
        page.setLayout(layout)
        return page

    def create_final_entry_screen(self):
        page = QWidget()
        layout = QVBoxLayout()
        
        self.btn_pick_single = QPushButton("Wybierz szablon")
        self.btn_pick_single.clicked.connect(self.pick_single_file)
        self.lbl_single_file = QLabel("<i>Nie wybrano szablonu</i>")
        layout.addWidget(self.btn_pick_single)
        layout.addWidget(self.lbl_single_file)
        
        self.btn_pick_final_path = QPushButton("Wybierz ścieżkę zapisu finalnego pliku")
        self.btn_pick_final_path.clicked.connect(self.pick_final_save_path)
        self.lbl_final_path = QLabel("<i>Nie wybrano lokalizacji zapisu</i>")
        layout.addWidget(self.btn_pick_final_path)
        layout.addWidget(self.lbl_final_path)
        
        layout.addStretch()
        
        nav_layout = QHBoxLayout()
        btn_back_2 = QPushButton("Wstecz")
        btn_back_2.clicked.connect(lambda: self.stack.setCurrentIndex(1))
        
        self.btn_print = QPushButton("Zatwierdź")
        self.btn_print.clicked.connect(self.full_report)
        
        nav_layout.addWidget(btn_back_2)
        nav_layout.addWidget(self.btn_print)
        layout.addLayout(nav_layout)
        page.setLayout(layout)
        
        self.validate_final_screen()
        return page

    def pick_folder(self):
        path = QFileDialog.getExistingDirectory(self, "Wybierz folder")
        if path:
            self.tree.clear()
            self.tree.blockSignals(True)
            self.create_tree_item(path, self.tree)
            self.tree.blockSignals(False)
            self.validate_screen1()

    def create_tree_item(self, path, parent):
        img_exts = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'}
        path_obj = pathlib.Path(path)
        is_dir = path_obj.is_dir()
        ext = path_obj.suffix.lower()
        
        if is_dir or ext in img_exts:
            item = QTreeWidgetItem(parent, [path_obj.name])
            item.setData(0, Qt.ItemDataRole.UserRole, str(path_obj))
            item.setCheckState(0, Qt.CheckState.Unchecked)
            
            try:
                if is_dir:
                    for entry in sorted(os.listdir(path)):
                        self.create_tree_item(os.path.join(path, entry), item)
                    if item.childCount() == 0:
                        # Remove empty folders from view
                        if isinstance(parent, QTreeWidget):
                            parent.takeTopLevelItem(parent.indexOfTopLevelItem(item))
                        else:
                            parent.removeChild(item)
            except PermissionError:
                pass

    def handle_check_state(self, item, column):
        self.tree.blockSignals(True)
        state = item.checkState(column)
        self.toggle_children(item, state)
        self.tree.blockSignals(False)
        self.validate_screen1()

    def toggle_children(self, item, state):
        for i in range(item.childCount()):
            child = item.child(i)
            child.setCheckState(0, state)
            self.toggle_children(child, state)

    def validate_screen1(self):
        any_checked = False
        it = QTreeWidgetItemIterator(self.tree)
        while it.value():
            item = it.value()
            path = item.data(0, Qt.ItemDataRole.UserRole)
            if item.checkState(0) == Qt.CheckState.Checked and os.path.isfile(path):
                any_checked = True
                break
            it += 1
        self.btn_next.setEnabled(any_checked)

    def go_to_preview(self):
        self.selected_files = []
        self.list_original.clear()
        it = QTreeWidgetItemIterator(self.tree)
        while it.value():
            item = it.value()
            path = item.data(0, Qt.ItemDataRole.UserRole)
            if item.checkState(0) == Qt.CheckState.Checked and os.path.isfile(path):
                self.selected_files.append(path)
                self.list_original.addItem(path)
            it += 1
        self.refresh_preview_list()
        self.stack.setCurrentIndex(1)

    def refresh_preview_list(self):
        self.list_converted.clear()
        suffix = self.name_suffix_input.text()
        for path in self.selected_files:
            self.list_converted.addItem(convert(path, suffix))

    def pick_single_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Wybierz plik szablonu")
        if path:
            self.manual_file_path = os.path.normpath(path)
            self.lbl_single_file.setText(f"Szablon: <i>{pathlib.Path(path).name}</i>")
        self.validate_final_screen()

    def pick_final_save_path(self):
        path, _ = QFileDialog.getSaveFileName(self, "Zapisz raport jako", "", "Word Documents (*.docx);;All Files (*)")
        if path:
            self.final_save_path = os.path.normpath(path)
            self.lbl_final_path.setText(f"Cel: <i>{pathlib.Path(path).name}</i>")
        self.validate_final_screen()

    def validate_final_screen(self):
        is_valid = bool(self.manual_file_path and self.final_save_path)
        self.btn_print.setEnabled(is_valid)

    def show_confirmation_popup(self):
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Sukces")
        msg_box.setText("Raport został utworzony")
        msg_box.setStandardButtons(QMessageBox.StandardButton.Ok)
        msg_box.exec()
        QApplication.quit()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = FileApp()
    window.show()
    sys.exit(app.exec())
